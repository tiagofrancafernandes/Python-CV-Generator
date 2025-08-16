from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Criar documento
doc = Document()

# --- Cabeçalho ---
nome = doc.add_paragraph("Tiago França Fernandes")
nome.style = doc.styles["Heading 1"]
nome.alignment = WD_ALIGN_PARAGRAPH.CENTER

funcao = doc.add_paragraph("Desenvolvedor Back-end Sênior | Especialista em PHP/Laravel & Infraestrutura")
funcao.alignment = WD_ALIGN_PARAGRAPH.CENTER

contatos = doc.add_paragraph("📧 tiago@tiagofranca.com | 📱 +55 41 98440-2684 (WhatsApp) | 🌐 tiagofranca.com\n"
                              "🔗 LinkedIn: linkedin.com/in/tiago-php | 💻 GitHub: github.com/tiagofrancafernandes")
contatos.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# --- Apresentação ---
doc.add_heading("Apresentação", level=2)
doc.add_paragraph(
    "Sou desenvolvedor back-end sênior com sólida experiência em PHP, Laravel, Linux e PostgreSQL, "
    "atuando desde 2009 na área de TI e desde 2013 com programação. "
    "Tenho histórico em projetos de alta performance, escalabilidade e segurança, atuando tanto no desenvolvimento "
    "quanto na arquitetura e infraestrutura de sistemas.\n\n"
    "Desde 2019, também atuo como professor particular de programação pelo Superprof Brasil, onde mantenho uma das "
    "melhores avaliações nacionais, ensinando desde fundamentos de programação até práticas avançadas como testes automatizados, "
    "DevOps e arquitetura de software. Meu foco é unir conhecimento técnico e clareza de comunicação para entregar resultados sólidos e sustentáveis."
)

# --- Competências Técnicas ---
doc.add_heading("Competências Técnicas", level=2)
skills = [
    "PHP, Laravel, Composer",
    "Vue.js, Nuxt, Alpine.js, Node.js, TypeScript",
    "PostgreSQL, MySQL/MariaDB, Redis",
    "AWS (EC2, RDS, SNS, S3, EKS, Lambda)",
    "Docker, Linux/Unix, Shell Script",
    "TDD, PHPUnit, Jest, OAuth/JWT",
    "Integração de APIs e microserviços",
    "SRE & DevOps mindset"
]
for skill in skills:
    doc.add_paragraph(f"• {skill}")

# --- Experiência Profissional ---
doc.add_heading("Experiência Profissional", level=2)
experiencias = [
    {
        "cargo": "Líder Técnico / Desenvolvedor Sênior",
        "empresa": "F2 Sistemas e Ensino",
        "periodo": "Nov 2023 - Atual",
        "descricao": [
            "Desenvolvimento de ferramentas internas e soluções SaaS para clientes.",
            "Aulas e treinamentos de programação.",
            "Projetos escaláveis com PHP, Laravel, Filament, Livewire, Vue.js e AWS."
        ]
    },
    {
        "cargo": "Desenvolvedor Back-end Sênior",
        "empresa": "Grupo GPS",
        "periodo": "Jun 2023 - Nov 2023",
        "descricao": [
            "Desenvolvimento e refatoração de sistemas centrais.",
            "Criação de APIs com PHP/Laravel e melhorias de infraestrutura."
        ]
    },
    {
        "cargo": "Desenvolvedor Back-end",
        "empresa": "Só Carrão",
        "periodo": "Jan 2023 - Mai 2023",
        "descricao": [
            "Desenvolvimento de APIs e aplicações PHP/Laravel.",
            "Integrações com AWS e otimização de performance."
        ]
    },
    {
        "cargo": "Sr. Back-end PHP Developer",
        "empresa": "Pontomais",
        "periodo": "Set 2019 - Out 2022",
        "descricao": [
            "Responsável pelo portal de parceiros externos e integrações de vendas.",
            "Implementação de recursos críticos de negócio e cálculos de comissão."
        ]
    },
    {
        "cargo": "Professor Particular de Programação",
        "empresa": "Superprof Brasil",
        "periodo": "2019 - Atual",
        "descricao": [
            "Aulas de programação do básico ao avançado.",
            "Conteúdo personalizado para PHP, Laravel, Docker, Linux e DevOps."
        ]
    }
]
for exp in experiencias:
    doc.add_paragraph(f"{exp['cargo']} - {exp['empresa']} ({exp['periodo']})").bold = True
    for desc in exp["descricao"]:
        doc.add_paragraph(f"• {desc}")

# --- Formação ---
doc.add_heading("Formação Acadêmica", level=2)
doc.add_paragraph("Sistemas para Internet - UniFCV (2021 - 2024)")
doc.add_paragraph("Sistemas de Informação - Estácio de Sá (2017 - 2021)")

# --- Certificações ---
doc.add_heading("Certificações e Cursos", level=2)
certs = [
    "PHP - Udemy, SENAI, PROLab",
    "Laravel - Udemy",
    "Linux/Unix - SENAI, Casa Brasil",
    "HTML5, CSS3, JavaScript - Udemy, PROLab Cursos",
    "Shell Script - Linux Cursos"
]
for cert in certs:
    doc.add_paragraph(f"• {cert}")

# --- Idiomas ---
doc.add_heading("Idiomas", level=2)
doc.add_paragraph("Português - Nativo")
doc.add_paragraph("Inglês - Básico/Intermediário (Leitura: 7/10, Escrita: 6/10, Conversação: 3/10)")

# Salvar arquivo
caminho_cv = "./output/CV-ptBR-Tiago-Franca-Fernandes.docx"
doc.save(caminho_cv)

caminho_cv
