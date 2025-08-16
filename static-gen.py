from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import os

def carregar_dados_curriculo(arquivo_json, idioma):
    with open(arquivo_json, 'r', encoding='utf-8') as f:
        dados = json.load(f)
    if idioma not in dados['idiomas']:
        raise ValueError(f"Idioma '{idioma}' não disponível no arquivo JSON.")
    return dados['curriculos'][idioma]

# Caminho do arquivo JSON e idioma desejado
ARQUIVO_JSON = './curriculum_data.json'
IDIOMA = 'pt'  # ou 'en' para inglês

dados = carregar_dados_curriculo(ARQUIVO_JSON, IDIOMA)

# Criar documento
doc = Document()

# --- Cabeçalho ---
nome = doc.add_paragraph(dados['nome'])
nome.style = doc.styles["Heading 1"]
nome.alignment = WD_ALIGN_PARAGRAPH.CENTER

funcao = doc.add_paragraph(dados['funcao'])
funcao.alignment = WD_ALIGN_PARAGRAPH.CENTER

contatos = doc.add_paragraph(dados['contatos'])
contatos.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")

# --- Apresentação ---
doc.add_heading(dados['apresentacao_titulo'], level=2)
doc.add_paragraph(dados['apresentacao'])

# --- Competências Técnicas ---
doc.add_heading(dados['competencias_titulo'], level=2)
for skill in dados['competencias']:
    doc.add_paragraph(f"• {skill}")

# --- Experiência Profissional ---
doc.add_heading(dados['experiencia_titulo'], level=2)
for exp in dados['experiencia']:
    doc.add_paragraph(f"{exp['cargo']} - {exp['empresa']} ({exp['periodo']})").bold = True
    for desc in exp["descricao"]:
        doc.add_paragraph(f"• {desc}")

# --- Formação ---
doc.add_heading(dados['formacao_titulo'], level=2)
for formacao in dados['formacao']:
    doc.add_paragraph(formacao)

# --- Certificações ---
doc.add_heading(dados['certificacoes_titulo'], level=2)
for cert in dados['certificacoes']:
    doc.add_paragraph(f"• {cert}")

# --- Idiomas ---
doc.add_heading(dados['idiomas_titulo'], level=2)
for idioma in dados['idiomas_lista']:
    doc.add_paragraph(idioma)

# Salvar arquivo
output_dir = './output'
os.makedirs(output_dir, exist_ok=True)
caminho_cv = f"{output_dir}/CV-{IDIOMA}-Tiago-Franca-Fernandes.docx"
doc.save(caminho_cv)

print(f"Currículo gerado: {caminho_cv}")
